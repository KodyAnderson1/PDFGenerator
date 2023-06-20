import os
from typing import List

import yaml
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import argparse
import platform
import subprocess
from dataclasses import dataclass

BAD_SPAN_TEXT = "Description of criterion"
IMG_WIDTH_INCHES = 6.5  # default image width in inches (can be changed). Default margins are 1 inch on each side


@dataclass
class SoupDetails:
    title: str
    sections: List[str]
    fileEnding: str


def read_yaml_file(file_path):
    with open(file_path) as yaml_file:
        data = yaml.safe_load(yaml_file)
    return data


def read_file(file_path):
    with open(file_path) as f:
        contents = f.read()
    return contents


def get_png_images(directory):
    return sorted([os.path.join(directory, f) for f in os.listdir(directory) if f.lower().endswith('.png')])


def create_word_document(data, images, soupDetails: SoupDetails):
    document = Document()

    # Add title and heading to the first page
    title = document.add_heading(soupDetails.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading = document.add_paragraph(data['heading']['Name'])
    heading.add_run(f"\n{data['heading']['Course']}")
    today = datetime.datetime.today().strftime('%B %d, %Y')
    heading.add_run(f"\n{today}")

    # Add each section to its own page
    for i, section in enumerate(soupDetails.sections):
        document.add_page_break()
        section_title = document.add_heading(section, level=1)
        section_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section_title.style.font.size = Pt(20)  # Add this line to set font size to 20
        section_page = section_title._element.xpath('w:pPr/w:pageBreakBefore')
        if section_page:
            section_page[0].set('w:val', 'true')

        if images and i < len(images):
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(images[i], width=Inches(IMG_WIDTH_INCHES))

    return document


def save_word_document(document_, output_file_path):
    document_.save(output_file_path)


def get_docx_file_name(data, output_directory, file_ending):
    name = data['heading']['StudentID']
    return os.path.join(output_directory, f"{name}_{file_ending}.docx")


def convert_to_pdf(input_file_path, pdf_directory_):
    # Ensure the file paths are in the correct format
    input_file_path = os.path.abspath(input_file_path)
    output_file_path = os.path.join(pdf_directory_, os.path.basename(input_file_path).replace('.docx', '.pdf'))
    print(f"Converting {input_file_path} to {output_file_path}")

    # Kill any existing Word processes
    # kill_word_processes()

    # Convert the Word document to PDF
    try:
        import docx2pdf
        docx2pdf.convert(input_file_path, output_file_path)
        print(f"Conversion complete. PDF saved at {output_file_path}")
    except ImportError:
        print("docx2pdf library is not installed. Please install it using 'pip install docx2pdf'.")


def kill_word_processes():
    if platform.system() == 'Windows':
        import psutil
        for proc in psutil.process_iter():
            try:
                if "WINWORD.EXE" == proc.name():
                    proc.kill()
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ProcessLookupError):
                pass
    elif platform.system() == 'Darwin':
        subprocess.call(['pkill', 'soffice'])
    else:
        print("Unsupported operating system. Process termination is not supported.")


def truncate_string(string, length=155):
    if len(string) > length:
        return string[:(length - 1)] + "..."
    else:
        return string


def get_soup_details(configuration_yaml) -> SoupDetails:
    rubric_files = os.path.join(configuration_yaml['rubricDirectory'], configuration_yaml['rubricFile'])
    soup_deets = SoupDetails(title="", sections=[], fileEnding="")

    raw_html = read_file(rubric_files)
    soup = BeautifulSoup(raw_html, 'html.parser')

    relevant_spans = soup.find_all('span', class_='description description_title')
    title = soup.find('title')

    if title:
        soup_deets.title = title.get_text(strip=True)
    else:
        soup_deets.title = configuration_yaml['rubricFile'].split('.')[0].replace(" ", "").lower()

    for span in relevant_spans:
        txt = span.get_text(strip=True)
        if txt != BAD_SPAN_TEXT:
            truncated_txt = truncate_string(txt, configuration_yaml['truncateLength'])
            soup_deets.sections.append(truncated_txt)

    soup_deets.fileEnding = configuration_yaml['rubricFile'].split('.')[0].replace(" ", "").lower()

    return soup_deets


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    config = read_yaml_file('input.yaml')

    images = get_png_images(config['imgDirectory'])
    docx_directory = config.get('docxDirectory', '.')
    pdf_directory = config.get('pdfDirectory', '.')

    soup_details = get_soup_details(config)

    document = create_word_document(config, images, soup_details)
    docx_file_path = get_docx_file_name(config, docx_directory, soup_details.fileEnding)
    save_word_document(document, docx_file_path)

    # Convert the Word document to a PDF
    input('Press enter to convert the Word document to a PDF or CTRL + C to exit.\n'
          'If you want to add a table of contents, you need to add it manually.\n'
          'MAKE SURE YOU CLOSE THE WORD DOCUMENT BEFORE YOU PRESS ENTER\n')
    convert_to_pdf(docx_file_path, pdf_directory)
