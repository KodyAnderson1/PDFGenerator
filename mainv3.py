import os
import yaml
import datetime
import argparse
from bs4 import BeautifulSoup
from typing import List
from dataclasses import dataclass
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMG_WIDTH_INCHES = 6.5
BAD_SPAN_TEXT = "Description of criterion"


@dataclass
class SoupDetails:
    title: str
    sections: List[str]
    fileEnding: str


def read_yaml_file(file_path: str) -> dict:
    with open(file_path) as yaml_file:
        data = yaml.safe_load(yaml_file)
    return data


def read_file(file_path: str) -> str:
    with open(file_path) as f:
        contents = f.read()
    return contents


def get_png_images(directory: str) -> List[str]:
    return sorted([os.path.join(directory, f) for f in os.listdir(directory) if f.lower().endswith('.png')])


def create_word_document(data: dict, images: List[str], soup_details: SoupDetails) -> Document:
    document = Document()
    add_title_and_heading(document, data, soup_details)
    add_sections_to_document(document, soup_details, images)
    return document


def add_title_and_heading(document: Document, data: dict, soup_details: SoupDetails):
    title = document.add_heading(soup_details.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading = document.add_paragraph(data['heading']['Name'])
    heading.add_run(f"\n{data['heading']['Course']}")
    today = datetime.datetime.today().strftime('%B %d, %Y')
    heading.add_run(f"\n{today}")


def add_sections_to_document(document: Document, soup_details: SoupDetails, images: List[str]):
    for i, section in enumerate(soup_details.sections):
        document.add_page_break()
        section_title = document.add_heading(section, level=1)
        section_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section_title.style.font.size = Pt(20)
        if images and i < len(images):
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(images[i], width=Inches(IMG_WIDTH_INCHES))


def save_word_document(document: Document, output_file_path: str):
    document.save(output_file_path)


def get_docx_file_name(data: dict, output_directory: str, file_ending: str) -> str:
    name = data['heading']['StudentID']
    return os.path.join(output_directory, f"{name}_{file_ending}.docx")


def convert_to_pdf(input_file_path: str, pdf_directory: str):
    input_file_path = os.path.abspath(input_file_path)
    output_file_path = os.path.join(pdf_directory, os.path.basename(input_file_path).replace('.docx', '.pdf'))
    try:
        import docx2pdf
        docx2pdf.convert(input_file_path, output_file_path)
    except ImportError:
        print("docx2pdf library is not installed. Please install it using 'pip install docx2pdf'.")


def truncate_string(string: str, length: int = 155) -> str:
    if len(string) > length:
        return string[:(length - 1)] + "..."
    else:
        return string


def get_soup_details(configuration_yaml: dict) -> SoupDetails:
    rubric_files = os.path.join(configuration_yaml['rubricDirectory'], configuration_yaml['rubricFile'])
    raw_html = read_file(rubric_files)
    soup = BeautifulSoup(raw_html, 'html.parser')
    relevant_spans = soup.find_all('span', class_='description description_title')
    title = soup.find('title').get_text(strip=True) if soup.find('title') else \
    configuration_yaml['rubricFile'].split('.')[0].lower()
    sections = [truncate_string(span.get_text(strip=True), configuration_yaml['truncateLength']) for span in
                relevant_spans if span.get_text(strip=True) != BAD_SPAN_TEXT]
    file_ending = configuration_yaml['rubricFile'].split('.')[0].lower()
    return SoupDetails(title, sections, file_ending)


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    config = read_yaml_file('input.yaml')
    images = get_png_images(config['imgDirectory'])
    docx_directory = config.get('docxDirectory', '.')
    pdf_directory = config.get('pdfDirectory', '.')
    soup_details = get_soup_details(config)
    document = create_word_document(config, images, soup_details)
    docx_file_path = get_docx_file_name(config, docx_directory, soup_details.fileEnding)
    save_word_document(document, docx_file_path)
    input('Press enter to convert the Word document to a PDF or CTRL + C to exit.\n'
          'If you want to add a table of contents, you need to add it manually.\n'
          'MAKE SURE YOU CLOSE THE WORD DOCUMENT BEFORE YOU PRESS ENTER\n')
    convert_to_pdf(docx_file_path, pdf_directory)
